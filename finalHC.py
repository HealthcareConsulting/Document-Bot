#!/usr/bin/env python3
# ndis_cli_hybrid_cover_logo_v2_5_version_control_enhanced.py
# Enhanced image-safe variant with improved textbox detection and smart logo sizing
# + Right-aligned logo positioning + Cover page logos for policy manuals
# + Version control table handling with auto-dates
# + Smart context-aware logo sizing for headers and textboxes
# + FIXED: Proper page breaks for version control tables

import argparse, csv, json, re, shutil, zipfile
from pathlib import Path
from typing import Dict, List, Set, Tuple, Optional
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from datetime import datetime
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


PLACEHOLDER_LOGO = "<logo>"

# ============================================================================
# VERSION CONTROL TABLE FUNCTIONS
# ============================================================================

def get_ordinal_date(date_obj: datetime) -> str:
    """Convert datetime to ordinal format like '14th of August 2025'"""
    day = date_obj.day
    month = date_obj.strftime("%B")  # Full month name
    year = date_obj.year
    
    # Add ordinal suffix
    if 10 <= day <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    
    return f"{day}{suffix} of {month} {year}"

def get_version_control_dates() -> Tuple[str, str]:
    """Get current date and next year date in ordinal format"""
    today = datetime.now()
    next_year = today.replace(year=today.year + 1)
    
    current_date = get_ordinal_date(today)
    next_review_date = get_ordinal_date(next_year)
    
    return current_date, next_review_date

def find_and_update_version_control_table(doc: Document) -> bool:
    """Find existing version control table and update dates in place while preserving formatting"""
    try:
        current_date, next_review_date = get_version_control_dates()
        print(f"Looking for version control table to update with dates: Current = {current_date}, Next Review = {next_review_date}")
        
        # Look for tables that contain version control data
        for table in doc.tables:
            if table.rows:
                # Check all text in the table to see if it's version control related
                table_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        table_text += cell.text.lower() + " "
                
                # If this looks like a version control table
                if any(keyword in table_text for keyword in ["drafted", "version control", "reviewed", "amendment"]):
                    print(f"Found version control table, updating dates...")
                    
                    # Update dates in the table while preserving formatting
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text
                            
                            if not text.strip():
                                continue
                            
                            # Look for date patterns and replace them
                            date_patterns = [
                                r'\d+(?:st|nd|rd|th)\s+of\s+\w+\s+\d{4}',  # "1st of June 2025"
                                r'\w+\s+\d{4}',  # "June 2025", "May 2025"
                                r'\d{4}'  # Just year "2025"
                            ]
                            
                            updated = False
                            new_text = text
                            
                            for pattern in date_patterns:
                                if re.search(pattern, text, re.IGNORECASE):
                                    # Check if this might be next review (look for 2026 or higher year)
                                    if re.search(r'202[6-9]|20[3-9]\d', text):
                                        # This is likely next review date
                                        new_text = re.sub(pattern, next_review_date, text, flags=re.IGNORECASE)
                                        print(f"Updated next review date: '{text}' -> '{new_text}'")
                                        updated = True
                                    elif re.search(r'202[0-5]', text):
                                        # This is likely current date
                                        new_text = re.sub(pattern, current_date, text, flags=re.IGNORECASE)
                                        print(f"Updated current date: '{text}' -> '{new_text}'")
                                        updated = True
                                    break
                            
                            # If we updated the text, preserve formatting while updating
                            if updated:
                                # Store formatting from the first run
                                original_formatting = {}
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    first_run = cell.paragraphs[0].runs[0]
                                    original_formatting = {
                                        'font_name': first_run.font.name,
                                        'font_size': first_run.font.size,
                                        'is_bold': first_run.font.bold,
                                        'is_italic': first_run.font.italic,
                                        'font_color': first_run.font.color.rgb if first_run.font.color else None
                                    }
                                
                                # Update the cell text
                                cell.text = new_text
                                
                                # Restore formatting to the new text
                                try:
                                    if cell.paragraphs and cell.paragraphs[0].runs:
                                        for run in cell.paragraphs[0].runs:
                                            if original_formatting.get('font_name'):
                                                run.font.name = original_formatting['font_name']
                                            if original_formatting.get('font_size'):
                                                run.font.size = original_formatting['font_size']
                                            if original_formatting.get('is_bold') is not None:
                                                run.font.bold = original_formatting['is_bold']
                                            if original_formatting.get('is_italic') is not None:
                                                run.font.italic = original_formatting['is_italic']
                                            if original_formatting.get('font_color'):
                                                run.font.color.rgb = original_formatting['font_color']
                                        print(f"Preserved formatting: font={original_formatting.get('font_name')}, bold={original_formatting.get('is_bold')}")
                                except Exception as fmt_error:
                                    print(f"Warning: Could not fully preserve formatting: {fmt_error}")
                    
                    return True
        
        print("No version control table found")
        return False
        
    except Exception as e:
        print(f"Error updating version control table: {e}")
        return False

# def move_version_control_to_own_page(doc: Document) -> bool:
#     """
#     Enhanced version: Move version control content to its own page using proper page breaks
#     Works reliably even with tables and complex layouts
#     """
#     try:
#         # Find version control heading
#         version_heading_index = -1
#         version_table = None
#         version_table_element = None
        
#         for i, paragraph in enumerate(doc.paragraphs):
#             text = paragraph.text.strip().lower()
#             if "version control" in text and "table" in text:
#                 version_heading_index = i
#                 print(f"Found version control heading at paragraph {i}: '{paragraph.text}'")
#                 break
        
#         # Find version control table and its element
#         for table in doc.tables:
#             if table.rows:
#                 table_text = ""
#                 for row in table.rows:
#                     for cell in row.cells:
#                         table_text += cell.text.lower() + " "
                
#                 if any(keyword in table_text for keyword in ["drafted", "version control", "reviewed", "amendment"]):
#                     version_table = table
#                     version_table_element = table._element
#                     print(f"Found version control table")
#                     break
        
#         if version_heading_index == -1 and version_table is None:
#             print("No version control content found to move")
#             return False
        
#         # METHOD 1: Insert proper page break before version control heading
#         if version_heading_index >= 0:
#             try:
#                 version_paragraph = doc.paragraphs[version_heading_index]
                
#                 # Insert a new paragraph before the version control heading
#                 new_para = version_paragraph.insert_paragraph_before()
                
#                 # Add a page break to the new paragraph
#                 run = new_para.add_run()
#                 #run.add_break(WD_BREAK.PAGE)
                
#                 print("Successfully inserted page break before version control heading")
#                 return True
                
#             except Exception as heading_error:
#                 print(f"Could not insert page break before heading: {heading_error}")




def move_version_control_to_own_page(doc: Document) -> bool:
    """
    Ensure Version Control heading and its table stay together on one page,
    remove blank paragraphs, and avoid page breaks.
    """
    try:
        # Find heading
        version_heading_index = -1
        for i, para in enumerate(doc.paragraphs):
            if "version control" in para.text.strip().lower() and "table" in para.text.strip().lower():
                version_heading_index = i
                break

        if version_heading_index == -1:
            print("No version control heading found")
            return False

        heading_para = doc.paragraphs[version_heading_index]

        # Remove empty paragraphs and page breaks immediately after heading
        next_index = version_heading_index + 1
        while next_index < len(doc.paragraphs) and not doc.paragraphs[next_index].text.strip():
            p = doc.paragraphs[next_index]
            p._element.getparent().remove(p._element)

        # Remove empty paragraphs before heading
        prev_index = version_heading_index - 1
        while prev_index >= 0 and not doc.paragraphs[prev_index].text.strip():
            p = doc.paragraphs[prev_index]
            p._element.getparent().remove(p._element)
            prev_index -= 1
        # Set keep_with_next for heading
        # Set heading spacing
        heading_para.paragraph_format.space_before = Pt(0)  # remove space before
        heading_para.paragraph_format.space_after = Pt(6)  # small space after
        heading_para.paragraph_format.keep_with_next = True

        # Find the table that comes after the heading
        version_table = None
        for table in doc.tables:
            # Compare table start index to heading index
            first_cell = table.cell(0,0)
            cell_para_index = doc.paragraphs.index(first_cell.paragraph)
            if cell_para_index > version_heading_index:
                version_table = table
                break

        if version_table:
            # Set first row to keep with next so table doesn't split
            for row in version_table.rows:
                for cell in row.cells:
                    cell_paras = cell.paragraphs
                    if cell_paras:
                        cell_paras[0].paragraph_format.keep_with_next = True
                break  # only first row

        print("Version Control heading and table will stay together")
        return True

    except Exception as e:
        print(f"Error processing Version Control: {e}")
        return False

  




        
        # METHOD 2: If no heading found, insert page break before version control table
        if version_table_element is not None:
            try:
                # Find the paragraph that contains or precedes this table
                parent = version_table_element.getparent()
                
                # Create a new paragraph element with page break
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls, qn
                
                # Create a paragraph with a page break
                page_break_xml = f'''
                <w:p {nsdecls('w')}>
                    <w:r>
                        <w:br w:type="page"/>
                    </w:r>
                </w:p>
                '''
                
                page_break_element = parse_xml(page_break_xml)
                
                # Insert the page break before the table
                parent.insert(parent.index(version_table_element), page_break_element)
                
                print("Successfully inserted page break before version control table")
                return True
                
            except Exception as table_error:
                print(f"Could not insert page break before table: {table_error}")
        
        #METHOD 3: Fallback - find any content with "version control" and add page break
        try:
            # Look through all paragraphs for version control content
            for i, paragraph in enumerate(doc.paragraphs):
                if "version control" in paragraph.text.lower():
                    # Insert page break before this paragraph
                    new_para = paragraph.insert_paragraph_before()
                    run = new_para.add_run()
                    run.add_break(WD_BREAK.PAGE)
                    print(f"Fallback: Inserted page break before paragraph {i}")
                    return True
        except Exception as fallback_error:
            print(f"Fallback method failed: {fallback_error}")
        
        # METHOD 4: Ultra-fallback using line breaks with enhanced calculation
        try:
            print("Using enhanced line break method as ultra-fallback...")
            
            # Calculate position more accurately
            content_before_vc = 0
            
            # Count all content before version control (paragraphs, tables, etc.)
            found_vc = False
            for i, paragraph in enumerate(doc.paragraphs):
                if "version control" in paragraph.text.lower():
                    content_before_vc = i
                    found_vc = True
                    break
            
            if not found_vc and version_heading_index >= 0:
                content_before_vc = version_heading_index
            
            # Count tables before version control position
            tables_before = 0
            for table in doc.tables:
                # Rough heuristic: if table appears early in document, count it
                table_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        table_text += cell.text.lower() + " "
                
                # If this is NOT the version control table and appears before it
                if not any(keyword in table_text for keyword in ["drafted", "version control", "reviewed", "amendment"]):
                    tables_before += 1
            
            # Enhanced line break calculation
            # Base breaks needed
            base_breaks = 0 #MADE 10 BY KAVITA
            
            # Add more breaks for tables (tables take up more space)
            table_adjustment = tables_before * 6  # Each table needs ~8 additional breaks
            
            # Add breaks based on content position
            position_adjustment = max(0, 20 - content_before_vc)  # More breaks if VC appears early
            
            total_breaks = base_breaks + table_adjustment + position_adjustment
            
            # Cap at reasonable maximum
            total_breaks = min(total_breaks, 50)
            
            print(f"Break calculation: base={base_breaks}, tables={tables_before}*8={table_adjustment}, position={position_adjustment}, total={total_breaks}")
            
            # Find insertion point
            if version_heading_index > 0:
                insert_point = doc.paragraphs[version_heading_index - 1]
            elif len(doc.paragraphs) > 0:
                insert_point = doc.paragraphs[-1]
            else:
                insert_point = doc.add_paragraph()
            
            # Add the calculated number of line breaks
            # for _ in range(total_breaks):
            #     insert_point.add_run().add_break()
            
            print(f"Enhanced fallback: Added {total_breaks} line breaks")
            return True
            
        except Exception as ultra_fallback_error:
            print(f"Ultra-fallback failed: {ultra_fallback_error}")
            return False
        
    except Exception as e:
        print(f"Error moving version control to own page: {e}")
        return False

def insert_page_break_before_element(element, doc: Document) -> bool:
    """
    Helper function to insert a proper page break before any document element
    """
    try:
        # Get the parent of the element
        parent = element.getparent()
        if parent is None:
            return False
        
        # Create page break XML
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        page_break_xml = f'''
        <w:p {nsdecls('w')}>
            <w:r>
                <w:br w:type="page"/>
            </w:r>
        </w:p>
        '''
        
        page_break_element = parse_xml(page_break_xml)
        
        # Insert before the target element
        element_index = list(parent).index(element)
        parent.insert(element_index, page_break_element)
        
        return True
        
    except Exception as e:
        print(f"Failed to insert page break: {e}")
        return False

def process_version_control_table(doc: Document) -> bool:
    """Main function to handle version control table processing"""
    try:
        # Step 1: Update dates in existing table
        updated = find_and_update_version_control_table(doc)
        
        # Step 2: Move to own page using proper page breaks
        moved = move_version_control_to_own_page(doc)
        
        return updated or moved
        
    except Exception as e:
        print(f"Error processing version control table: {e}")
        return False

# ============================================================================
# COVER PAGE FUNCTIONS
# ============================================================================

def add_cover_page_logo_large(doc: Document, logo: Optional[Path] = None, width_mm: float = 40.0) -> bool:
    """
    Add cover page logo positioned below banner but above Contents - DEBUGGING VERSION
    """
    if not logo or not logo.exists():
        print(f"Cover logo failed: Logo file missing or doesn't exist")
        print(f"Logo path: {logo}")
        print(f"Logo exists: {logo.exists() if logo else 'No logo provided'}")
        return False
    
    print(f"ATTEMPTING to insert cover logo at {width_mm}mm...")
    print(f"Logo file: {logo}")
    print(f"Logo size: {logo.stat().st_size} bytes")
    print(f"Document has {len(doc.paragraphs)} paragraphs")
    
    try:
        # METHOD 1: Insert at very beginning with detailed tracking
        print(f"Creating new paragraph for cover logo...")
        
        # Create paragraph FIRST
        new_p = doc.add_paragraph()
        print(f"New paragraph created")
        
        # Set alignment BEFORE moving
        new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        print(f"Alignment set to LEFT")
        
        # Add spacing BEFORE moving
        paragraph_format = new_p.paragraph_format
        paragraph_format.space_before = Mm(15)  # Reduced spacing
        paragraph_format.space_after = Mm(3)
        print(f"Spacing set: before={15}mm, after={3}mm")
        
        # Add logo BEFORE moving paragraph
        print(f"Adding logo to paragraph...")
        run = new_p.add_run()
        run.add_picture(str(logo), width=Mm(width_mm))
        print(f"Logo added to run with width {width_mm}mm")
        
        # Add line break after logo (no debug text)
        #run.add_break()
        print(f"Line break added after logo")
        
        # NOW move to beginning
        print(f"Moving paragraph to beginning...")
        new_p_element = new_p._element
        parent = new_p_element.getparent()
        parent.remove(new_p_element)
        parent.insert(0, new_p_element)
        print(f"Paragraph moved to position 0")
        
        # Verify placement
        print(f"Document now has {len(doc.paragraphs)} paragraphs")
        if doc.paragraphs:
            first_para_text = doc.paragraphs[0].text
            print(f"First paragraph text: '{first_para_text[:100]}...'")
        
        print(f"COVER LOGO INSERTION COMPLETED SUCCESSFULLY")
        return True
        
    except Exception as e:
        print(f"Cover logo insertion failed with error: {e}")
        import traceback
        print(f"Full traceback:")
        traceback.print_exc()
        return False

def is_policy_manual(doc_path: Path) -> bool:
    """
    Check if document should receive a cover logo based on filename patterns
    """
    filename_lower = doc_path.name.lower()
    
    # List of document types that should get cover logos
    cover_logo_patterns = [
        "policy and procedure manual",
        "business plan",
        "00",  # Matches documents with "00" in filename
        "policy and procedures",
        "handbook",
        "psychological assessment form",
        "risk assessment guide and checklist",
        "service agreement and schedule of support",
        "evaluation of competency"
    ]
    
    # Check if any pattern matches
    for pattern in cover_logo_patterns:
        if pattern in filename_lower:
            print(f"Cover logo check: '{filename_lower}' contains '{pattern}': True")
            return True
    
    print(f"Cover logo check: '{filename_lower}' - no matching patterns found: False")
    return False

# ============================================================================
# CORE REPLACEMENT FUNCTIONS
# ============================================================================

def load_replacements(p: Path) -> Dict[str, str]:
    data = json.loads(p.read_text(encoding="utf-8"))
    norm = {}
    for k, v in data.items():
        k = str(k)
        norm[k] = "" if v is None else str(v)
        norm[k.lower()] = "" if v is None else str(v)
    return norm

def discover_placeholders(text: str) -> Set[str]:
    return set(re.findall(r"<[^<>]+>", text or ""))

def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for t2 in cell.tables:
                    for r2 in t2.rows:
                        for c2 in r2.cells:
                            for p2 in c2.paragraphs:
                                yield p2

def iter_header_footer_paragraphs(doc: Document):
    for s in doc.sections:
        hdr = s.header
        for p in hdr.paragraphs:
            yield p
        for t in hdr.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        yield p
        ftr = s.footer
        for p in ftr.paragraphs:
            yield p
        for t in ftr.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        yield p

def replace_in_run_text(run_text: str, repl: Dict[str, str]):
    if not run_text:
        return run_text, False, False, set()
    text = run_text
    changed = False
    
    # Handle regular placeholder replacements
    for k, v in repl.items():
        if not (k.startswith("<") and k.endswith(">")):
            continue
        if k.lower() == PLACEHOLDER_LOGO:
            continue
        new = re.sub(re.escape(k), v, text, flags=re.IGNORECASE)
        if new != text:
            text = new
            changed = True
    
    # Handle smart possessive for company names ending in 's'
    text, possessive_changed = handle_smart_possessive(text, repl)
    if possessive_changed:
        changed = True
    
    logo_here = False
    if re.search(re.escape(PLACEHOLDER_LOGO), text, flags=re.IGNORECASE):
        logo_here = True
        text = re.sub(re.escape(PLACEHOLDER_LOGO), "", text, flags=re.IGNORECASE)
        changed = True
    unresolved = {tok for tok in discover_placeholders(text) if tok.strip().lower() != PLACEHOLDER_LOGO}
    return text, changed, logo_here, unresolved

def handle_smart_possessive(text: str, repl: Dict[str, str]) -> Tuple[str, bool]:
    """
    Handle smart possessive for company names ending in 's'
    Converts "<company name>'s" to "Support Services'" if company name ends with 's'
    """
    changed = False
    
    # Look for <company name>'s pattern
    possessive_pattern = r'<company\s+name>\s*\'s\b'
    
    if re.search(possessive_pattern, text, re.IGNORECASE):
        # Get the company name value
        company_name = repl.get("<company name>", repl.get("<company name>", ""))
        
        if company_name and company_name.strip().endswith('s'):
            # Replace <company name>'s with "Company Name'" (no extra 's')
            new_text = re.sub(possessive_pattern, f"{company_name}'", text, flags=re.IGNORECASE)
            print(f"Smart possessive: '{company_name}' ends with 's', using '{company_name}'' instead of '{company_name}'s'")
            return new_text, True
        else:
            # Normal possessive handling for non-s ending names
            if company_name:
                new_text = re.sub(possessive_pattern, f"{company_name}'s", text, flags=re.IGNORECASE)
                return new_text, True
    
    return text, changed

# ============================================================================
# ENHANCED LOGO INSERTION WITH SMART CONTEXT DETECTION
# ============================================================================

def detect_header_context(paragraph):
    """Detect if we're in a header and what type of context"""
    context = {
        'is_header': False,
        'is_footer': False,
        'is_in_table': False,
        'recommended_width': 35.0  # default
    }
    
    # Check if in header/footer
    try:
        part_type = paragraph._part.content_type
        if 'header' in part_type:
            context['is_header'] = True
            context['recommended_width'] = 20.0  # Much smaller for headers
        elif 'footer' in part_type:
            context['is_footer'] = True
            context['recommended_width'] = 15.0  # Even smaller for footers
    except:
        pass
    
    # Check if in table (common in headers)
    try:
        parent = paragraph._element.getparent()
        while parent is not None:
            if parent.tag.endswith('tc'):  # table cell
                context['is_in_table'] = True
                if context['is_header']:
                    context['recommended_width'] = 18.0  # Very small for header tables
                else:
                    context['recommended_width'] = 25.0  # Medium for body tables
                break
            parent = parent.getparent()
    except:
        pass
    
    return context

def process_par_safe_logo_smart(paragraph, repl: Dict[str, str], logo: Optional[Path] = None, width_mm: float = 35.0, dry: bool = False):
    """
    Enhanced safe version with smart logo sizing based on context
    """
    changed = False
    logo_trig = False
    unresolved_all: Set[str] = set()
    
    # Process text replacements
    for run in paragraph.runs:
        new, chg, l_here, unres = replace_in_run_text(run.text, repl)
        run.text = new
        if chg:
            changed = True
        if l_here:
            logo_trig = True
        unresolved_all.update(unres)
    
    logo_inserted = False
    if logo_trig and logo and not dry:
        
        # SMART SIZING: Detect context and adjust logo size
        context = detect_header_context(paragraph)
        smart_width = context['recommended_width']
        
        print(f"Logo context: header={context['is_header']}, table={context['is_in_table']}, width={smart_width}mm")
        
        # Set paragraph alignment to right
        try:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            print(f"Set paragraph alignment to right")
        except Exception as align_error:
            print(f"Could not set alignment: {align_error}")
        
        # Add the logo with context-appropriate size
        try:
            r = paragraph.add_run()
            r.add_picture(str(logo), width=Mm(smart_width))
            print(f"Logo inserted and right-aligned (width: {smart_width}mm)")
            logo_inserted = True
            changed = True
        except Exception as logo_error:
            print(f"Logo insertion failed: {logo_error}")
    
    return changed, logo_inserted, unresolved_all

# ============================================================================
# ENHANCED TEXTBOX PROCESSING
# ============================================================================

def process_shape_textboxes_enhanced(doc: Document, repl: Dict[str, str]) -> bool:
    """
    Enhanced processing of shape-based textboxes with better detection
    """
    changed = False
    shapes_processed = 0
    
    try:
        print(f"Processing shape-based textboxes for placeholders...")
        
        # Get the document XML root
        doc_element = doc.element
        
        # Define namespaces for drawing shapes
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        # Look for drawing elements that contain text
        drawings = doc_element.xpath('.//w:drawing', namespaces=namespaces)
        print(f"Found {len(drawings)} drawing elements")
        
        for i, drawing in enumerate(drawings):
            # Look for text content within drawing shapes
            try:
                # Method 1: Look for drawing paragraphs (a:p)
                drawing_paragraphs = drawing.xpath('.//a:p', namespaces=namespaces)
                if drawing_paragraphs:
                    print(f"Shape {i+1}: Found {len(drawing_paragraphs)} drawing paragraphs")
                    shapes_processed += 1
                    
                    for j, dp in enumerate(drawing_paragraphs):
                        # Get text runs within the drawing paragraph
                        text_runs = dp.xpath('.//a:t', namespaces=namespaces)
                        
                        for k, text_run in enumerate(text_runs):
                            original_text = text_run.text or ""
                            
                            if original_text.strip():
                                print(f"    Text in shape {i+1}, para {j+1}, run {k+1}: '{original_text}'")
                            
                            if "<" in original_text:
                                new_text = original_text
                                
                                # Replace placeholders (including ALL placeholders)
                                for placeholder, value in repl.items():
                                    if placeholder in new_text:
                                        new_text = new_text.replace(placeholder, value or "")
                                        changed = True
                                        print(f"    Replaced '{placeholder}' with '{value}' in shape {i+1}")
                                
                                # Handle smart possessive for shape text
                                new_text, possessive_changed = handle_smart_possessive(new_text, repl)
                                if possessive_changed:
                                    changed = True
                                
                                # Update the text
                                text_run.text = new_text
                
                # Method 2: Look for text in textboxes (w:txbxContent)
                textboxes = drawing.xpath('.//w:txbxContent', namespaces=namespaces)
                if textboxes:
                    print(f"Shape {i+1}: Found {len(textboxes)} textboxes")
                    for tb_idx, textbox in enumerate(textboxes):
                        tb_changed = process_textbox_content_enhanced(textbox, repl, f"{i+1}-{tb_idx+1}", namespaces)
                        if tb_changed:
                            changed = True
                            shapes_processed += 1
                
            except Exception as shape_error:
                print(f"Error processing shape {i+1}: {shape_error}")
        
        print(f"Processed {shapes_processed} shape-based textboxes")
        return changed
        
    except Exception as e:
        print(f"Error processing shape textboxes: {e}")
        import traceback
        traceback.print_exc()
        return False

def process_textbox_content_enhanced(textbox_element, repl: Dict[str, str], textbox_id: str, namespaces: dict) -> bool:
    """
    Enhanced processing of content within a single textbox element
    """
    changed = False
    
    try:
        # Get all text nodes in this textbox using proper namespaces
        text_nodes = textbox_element.xpath('.//w:t', namespaces=namespaces)
        print(f"Textbox {textbox_id}: Found {len(text_nodes)} text nodes")
        
        for i, text_node in enumerate(text_nodes):
            original_text = text_node.text or ""
            
            if original_text.strip():  # Only log non-empty text
                print(f"    Text node {i}: '{original_text}'")
            
            if "<" in original_text:
                new_text = original_text
                
                # Replace ALL placeholders (including logo placeholders)
                for placeholder, value in repl.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, value or "")
                        changed = True
                        print(f"    Replaced '{placeholder}' with '{value}' in textbox {textbox_id}")
                
                # Handle smart possessive for textbox content
                new_text, possessive_changed = handle_smart_possessive(new_text, repl)
                if possessive_changed:
                    changed = True
                
                # Update the text node
                text_node.text = new_text
        
        return changed
        
    except Exception as e:
        print(f"Error processing textbox content for {textbox_id}: {e}")
        return False

# ============================================================================
# XML PROCESSING FUNCTIONS (UNCHANGED BUT IMPROVED)
# ============================================================================

def _cross_run_replace_xml(p_elm, repl: Dict[str, str]):
    """Enhanced XML replacement with smart possessive handling"""
    t_nodes = p_elm.xpath('.//w:t')
    if not t_nodes:
        return False, False, set()
    texts = [(t.text or "") for t in t_nodes]
    changed = False
    logo = False
    tokens = [k for k in repl.keys() if k.startswith("<") and k.endswith(">")]
    tokens = list(dict.fromkeys(sorted(tokens, key=lambda s: -len(s))))

    def find(token_lower: str, sn: int = 0, so: int = 0):
        i = sn
        j = 0
        off = so
        while i < len(texts):
            nt = texts[i]
            while off < len(nt) and j < len(token_lower):
                if nt[off].lower() != token_lower[j]:
                    break
                off += 1
                j += 1
            if j == len(token_lower):
                return sn, so, i, off
            if off < len(nt) and nt[off].lower() != token_lower[j]:
                so += 1
                if so >= len(texts[sn]):
                    sn += 1
                    so = 0
                i = sn
                j = 0
                off = so
                continue
            i += 1
            off = 0
        return None

    # Join all text for possessive processing
    full_text = "".join(texts)
    
    # Handle smart possessive first
    full_text, possessive_changed = handle_smart_possessive(full_text, repl)
    if possessive_changed:
        changed = True
        # Update the texts array with the new text
        texts = [full_text]
        # Clear other nodes since we merged everything
        for i in range(1, len(t_nodes)):
            texts.append("")

    pn = 0
    po = 0
    while pn < len(texts):
        if all('<' not in t for t in texts[pn:]):
            break
        cur = texts[pn]
        if po < len(cur) and cur[po] != '<':
            po += 1
            if po >= len(cur):
                pn += 1
                po = 0
            continue
        matched = False
        for tok in tokens:
            tl = tok.lower()
            hit = find(tl, pn, po)
            if not hit:
                continue
            sn, so, en, eo = hit
            val = repl.get(tok, repl.get(tok.lower(), ""))
            if tl == PLACEHOLDER_LOGO:
                val = ""
                logo = True
            left = texts[sn][:so]
            right = texts[en][eo:]
            texts[sn] = left + str(val) + (right if sn == en else "")
            for k in range(sn + 1, en + 1):
                if k == en:
                    texts[k] = texts[k][eo:]
                else:
                    texts[k] = ""
            changed = True
            pn = sn
            po = so + len(str(val))
            matched = True
            break
        if not matched:
            po += 1
            if po >= len(cur):
                pn += 1
                po = 0
    for t, nt in zip(t_nodes, texts):
        t.text = nt
    after = "".join(texts)
    unresolved = {t for t in discover_placeholders(after) if t.strip().lower() != PLACEHOLDER_LOGO}
    return changed, logo, unresolved

def _has_image_anywhere_in(elt) -> bool:
    # Check for various image-related elements separately to avoid XPath complexity
    image_checks = [
        './/w:drawing',
        './/w:pict', 
        './/v:imagedata',
        './/a:blip',
        './/pic:pic',
        './/w:object'
    ]
    
    for xpath in image_checks:
        try:
            if elt.xpath(xpath):
                return True
        except:
            continue
    
    # Fallback: check XML content for image indicators
    try:
        xml_str = str(elt) if hasattr(elt, '__str__') else ""
        image_indicators = ['drawing', 'pict', 'imagedata', 'blip', 'graphic', 'pic:', 'embed']
        return any(indicator in xml_str.lower() for indicator in image_indicators)
    except:
        return False

def _ancestor_with_images(p_elm) -> bool:
    # Simplified ancestor checking - check each ancestor type separately
    ancestor_types = ['tc', 'tr', 'tbl', 'sdt', 'txbxContent', 'hdr', 'ftr']
    
    for anc_type in ancestor_types:
        try:
            ancestors = p_elm.xpath(f'ancestor::*[local-name()="{anc_type}"]')
            for ancestor in ancestors:
                if _has_image_anywhere_in(ancestor):
                    return True
        except:
            continue
    
    return False

def _paragraph_contains_image(p_elm) -> bool:
    """Check if paragraph itself directly contains image content"""
    # Simplified image detection in paragraph runs
    try:
        # Check for drawing elements
        if p_elm.xpath('.//w:drawing'):
            return True
        # Check for picture elements  
        if p_elm.xpath('.//w:pict'):
            return True
        # Check for objects
        if p_elm.xpath('.//w:object'):
            return True
    except:
        pass
    
    # Fallback: string-based detection
    try:
        xml_content = str(p_elm) if hasattr(p_elm, '__str__') else ""
        return any(indicator in xml_content.lower() for indicator in ['<w:drawing', '<w:pict', '<w:object', 'imagedata', 'blip'])
    except:
        return False

def prune_or_rescue_body_shapes(doc: Document, repl: Dict[str, str], report_missing: Set[str]) -> Tuple[int, int, int]:
    changed_count = 0
    logo_hits = 0
    pruned_count = 0

    root = doc.element.body
    paragraphs = list(root.xpath('.//w:p'))
    token_re = re.compile(r"<[^<>]+>")

    # Enhanced image presence detection
    has_img_here = [_has_image_anywhere_in(p) or _paragraph_contains_image(p) for p in paragraphs]

    for idx, p in enumerate(paragraphs):
        joined = "".join((t.text or "") for t in p.xpath('.//w:t'))
        jl = joined.lower()

        # Never prune the paragraph that has <logo>
        if "<logo>" in jl:
            if "<" in joined:
                chg, lg, unresolved = _cross_run_replace_xml(p, repl)
                if chg: changed_count += 1
                if lg:  logo_hits += 1
                report_missing.update(unresolved)
            continue

        # ENHANCED: Skip any paragraph that directly contains images
        if _paragraph_contains_image(p):
            # Still do replacements but never prune
            if "<" in joined:
                chg, lg, unresolved = _cross_run_replace_xml(p, repl)
                if chg: changed_count += 1
                if lg:  logo_hits += 1
                report_missing.update(unresolved)
            continue

        # Rescue split placeholders first
        chg, lg, unresolved = _cross_run_replace_xml(p, repl)
        if chg: changed_count += 1
        if lg:  logo_hits += 1
        report_missing.update(unresolved)

        # Re-read after rescue
        joined = "".join((t.text or "") for t in p.xpath('.//w:t'))
        jl = joined.lower()

        tokens_in_par = [tok.lower() for tok in token_re.findall(jl)]
        if not tokens_in_par:
            continue

        # Ultra image-protection with enhanced detection:
        neighbour_image = (
            has_img_here[idx] or 
            (idx > 0 and has_img_here[idx-1]) or 
            (idx+1 < len(has_img_here) and has_img_here[idx+1])
        )
        container_image = _ancestor_with_images(p)
        
        # ADDITIONAL SAFETY: Check for any image-related content more safely
        try:
            has_image_indicators = _has_image_anywhere_in(p) or _paragraph_contains_image(p)
        except:
            has_image_indicators = True  # If we can't check safely, assume there are images
        
        if neighbour_image or container_image or has_image_indicators:
            # Do not prune anything near images or with image indicators
            continue

        # Additional safety: Check if paragraph is very short (likely contains only an image)
        text_content = joined.strip()
        if len(text_content) < 50 and not tokens_in_par:  # Very short, no placeholders
            continue

        # Remove only if ALL tokens are blank/missing AND no image protection applies
        all_blank = True
        for tok in tokens_in_par:
            val = repl.get(tok, repl.get(tok.lower(), None))
            if val is not None and str(val).strip() != "":
                all_blank = False
                break
        
        if all_blank:
            # Final safety check before removal
            if not _has_image_anywhere_in(p) and not _paragraph_contains_image(p):
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)
                    pruned_count += 1

    return changed_count, logo_hits, pruned_count

def rescue_header_footer_shapes(doc: Document, repl: Dict[str, str], report_missing: Set[str]) -> Tuple[int, int]:
    changed_count = 0
    logo_hits = 0
    for part in doc.part.package.parts:
        if part.content_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml",
        ):
            for p in part.element.xpath('.//w:p'):
                joined = "".join((t.text or "") for t in p.xpath('.//w:t'))
                if "<" not in joined:
                    continue
                
                # Enhanced: Don't modify paragraphs with images in headers/footers
                if _has_image_anywhere_in(p):
                    continue
                    
                chg, lg, unresolved = _cross_run_replace_xml(p, repl)
                if chg:
                    changed_count += 1
                if lg:
                    logo_hits += 1
                report_missing.update(unresolved)
    return changed_count, logo_hits

# ============================================================================
# MAIN DOCUMENT PROCESSING
# ============================================================================

def process_headers_and_footers_original_safe(doc: Document, repl: Dict[str, str], logo: Optional[Path] = None, width_mm: float = 15.0, dry: bool = False):
    """
    ORIGINAL SAFE processing of headers and footers - NO alignment changes
    """
    changed = False
    logos_inserted = 0
    
    try:
        print(f"Processing headers and footers (ORIGINAL SAFE method)...")
        
        for section in doc.sections:
            # Process header
            if section.header:
                header_changed, header_logos = process_header_footer_content_original_safe(
                    section.header, repl, logo, width_mm, dry, "Header"
                )
                if header_changed:
                    changed = True
                logos_inserted += header_logos
            
            # Process footer
            if section.footer:
                footer_changed, footer_logos = process_header_footer_content_original_safe(
                    section.footer, repl, logo, width_mm, dry, "Footer"
                )
                if footer_changed:
                    changed = True
                logos_inserted += footer_logos
        
        print(f"Processed headers/footers ORIGINAL SAFE method, inserted {logos_inserted} logos")
        return changed, logos_inserted
        
    except Exception as e:
        print(f"Error processing headers/footers: {e}")
        return False, 0

def process_header_footer_content_original_safe(hf_part, repl: Dict[str, str], logo: Optional[Path] = None, width_mm: float = 15.0, dry: bool = False, part_type: str = "Header"):
    """
    ORIGINAL SAFE processing - simple text replacement and logo insertion WITHOUT alignment changes
    """
    changed = False
    logos_inserted = 0
    
    try:
        # Process all paragraphs in header/footer
        for paragraph in hf_part.paragraphs:
            para_changed, para_logos = process_header_paragraph_original_safe(
                paragraph, repl, logo, width_mm, dry, part_type
            )
            if para_changed:
                changed = True
            if para_logos:
                logos_inserted += para_logos
        
        # Process tables in header/footer
        for table in hf_part.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_changed, para_logos = process_header_paragraph_original_safe(
                            paragraph, repl, logo, width_mm, dry, f"{part_type} Table"
                        )
                        if para_changed:
                            changed = True
                        if para_logos:
                            logos_inserted += para_logos
        
        return changed, logos_inserted
        
    except Exception as e:
        print(f"Error processing {part_type}: {e}")
        return False, 0

def process_header_paragraph_original_safe(paragraph, repl: Dict[str, str], logo: Optional[Path] = None, width_mm: float = 15.0, dry: bool = False, context: str = "Header"):
    """
    ORIGINAL SAFE processing - NO alignment changes, simple logo insertion
    """
    changed = False
    logo_found = False
    
    try:
        # First pass: text replacement and logo detection
        for run in paragraph.runs:
            original_text = run.text
            if original_text and "<" in original_text:
                new_text = original_text
                
                # Replace all placeholders except <logo>
                for placeholder, value in repl.items():
                    if placeholder.lower() != "<logo>" and placeholder in new_text:
                        new_text = new_text.replace(placeholder, value or "")
                        changed = True
                
                # Check for logo placeholder
                if "<logo>" in new_text.lower():
                    logo_found = True
                    # Remove logo placeholder text
                    new_text = re.sub(r"<logo>", "", new_text, flags=re.IGNORECASE)
                    changed = True
                
                run.text = new_text
        
        # Second pass: insert logo if found (NO ALIGNMENT CHANGES)
        logos_inserted = 0
        if logo_found and logo and logo.exists() and not dry:
            try:
                print(f"Inserting logo in {context} (width: {width_mm}mm) - NO alignment changes")
                
                # Simple insertion with NO alignment changes (like original safe version)
                run = paragraph.add_run()
                run.add_picture(str(logo), width=Mm(width_mm))
                
                logos_inserted = 1
                changed = True
                print(f"Logo inserted successfully in {context} (no alignment changes)")
                
            except Exception as logo_error:
                print(f"Logo insertion failed in {context}: {logo_error}")
        
        return changed, logos_inserted
        
    except Exception as e:
        print(f"Error processing paragraph in {context}: {e}")
        return False, 0

def process_docx(input_path: Path, output_path: Path, repl: Dict[str, str], logo: Optional[Path] = None, width_mm: float = 35.0, dry: bool = False):
    """
    Enhanced processing with smart logo sizing and comprehensive textbox handling
    """
    print(f"Processing: {input_path.name}")
    
    report = {
        "file": str(input_path),
        "output": str(output_path),
        "changed": False,
        "placeholders_found": set(),
        "placeholders_missing": set(),
        "logos_inserted_body": 0,
        "logos_inserted_headers": 0,
        "version_control_processed": 0,
        "xml_logo_hits": 0,
        "xml_paras_changed": 0,
        "xml_paras_pruned": 0,
    }
    
    try:
        doc = Document(str(input_path))
        print(f"Document loaded successfully")
    except Exception as e:
        print(f"Failed to load document: {e}")
        return report

    # 1. Add cover logo for policy manuals (ENHANCED DEBUGGING)
    try:
        print(f"=== COVER LOGO DEBUG START ===")
        print(f"Document name: {input_path.name}")
        print(f"Full path: {input_path}")
        
        is_policy = is_policy_manual(input_path)
        print(f"Is policy manual: {is_policy}")
        print(f"Logo provided: {logo is not None}")
        print(f"Logo exists: {logo and logo.exists() if logo else False}")
        print(f"Dry run: {dry}")
        print(f"Width setting: {width_mm}mm")
        
        # CHECK ALL CONDITIONS
        if not is_policy:
            print(f"SKIPPING: Not a policy manual")
        elif not logo:
            print(f"SKIPPING: No logo provided")
        elif not logo.exists():
            print(f"SKIPPING: Logo file doesn't exist at {logo}")
        elif dry:
            print(f"SKIPPING: Dry run mode")
        else:
            print(f"ALL CONDITIONS MET - PROCEEDING WITH COVER LOGO")
            
            # Document state before
            print(f"Document state BEFORE logo insertion:")
            print(f"    - Paragraphs: {len(doc.paragraphs)}")
            print(f"    - Tables: {len(doc.tables)}")
            
            # ATTEMPT INSERTION
            if add_cover_page_logo_large(doc, logo, width_mm):
                report["changed"] = True
                print(f"COVER LOGO FUNCTION RETURNED TRUE")
                
                # Document state after
                print(f"Document state AFTER logo insertion:")
                print(f"    - Paragraphs: {len(doc.paragraphs)}")
                print(f"    - Tables: {len(doc.tables)}")
                if doc.paragraphs:
                    print(f"    - First paragraph text: '{doc.paragraphs[0].text[:50]}...'")
            else:
                print(f"COVER LOGO FUNCTION RETURNED FALSE")
        
        print(f"=== COVER LOGO DEBUG END ===")
            
    except Exception as cover_error:
        print(f"Error in cover logo section: {cover_error}")
        import traceback
        traceback.print_exc()

    # 2. Process shape-based textboxes (ENHANCED - for cover page placeholders in shapes)
    try:
        print(f"Processing shape-based textboxes...")
        shape_changed = process_shape_textboxes_enhanced(doc, repl)
        if shape_changed:
            report["changed"] = True
            print(f"Shape textboxes processed successfully")
    except Exception as shape_error:
        print(f"Error processing shape textboxes: {shape_error}")

    # 3. Process body paragraphs - PROTECT COVER LOGO PARAGRAPH
    try:
        paragraph_count = 0
        is_policy = is_policy_manual(input_path)
        
        print(f"=== BODY PROCESSING DEBUG START ===")
        print(f"Is policy manual: {is_policy}")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
        for i, p in enumerate(doc.paragraphs):
            paragraph_count += 1
            para_text = p.text[:100] if p.text else "[EMPTY]"
            print(f"Processing paragraph {i}: '{para_text}...'")
            
            # STRICT PROTECTION: Skip first paragraph for policy manuals
            if is_policy and i == 0:
                print(f"PROTECTING first paragraph in policy manual (contains cover logo)")
                continue
            
            # Also skip any paragraph that has images (safer protection)
            has_image = hasattr(p, '_element') and _paragraph_contains_image(p._element)
            if has_image:
                print(f"PROTECTING paragraph {i} (contains images)")
                continue
            
            # Use the smart logo function for other paragraphs
            chg, logo_ins, unres = process_par_safe_logo_smart(p, repl, logo=logo, width_mm=width_mm, dry=dry)
            if chg:
                report["changed"] = True
                print(f"Paragraph {i} changed")
            if logo_ins:
                report["logos_inserted_body"] += logo_ins
                print(f"Logo inserted in paragraph {i}")
            report["placeholders_missing"].update(unres)
        
        print(f"Processed {paragraph_count} paragraphs (protected: {1 if is_policy else 0})")
        print(f"=== BODY PROCESSING DEBUG END ===")
        
    except Exception as process_error:
        print(f"Error during body processing: {process_error}")

    # 4. Process headers and footers (ORIGINAL SAFE METHOD - no alignment changes)
    try:
        # Use smaller size for headers (cap at 20mm) - separate from cover logos
        header_width = min(width_mm, 20.0)  # Cap at 20mm for headers only
        print(f"Header logo size: {header_width}mm (capped at 20mm for headers)")
        
        hf_changed, hf_logos = process_headers_and_footers_original_safe(doc, repl, logo, header_width, dry)
        if hf_changed:
            report["changed"] = True
        report["logos_inserted_headers"] = hf_logos
        
    except Exception as hf_error:
        print(f"Error processing headers/footers: {hf_error}")

    # 5. Shapes + XML pruning (ultra image protected)
    try:
        xml_changed, xml_logo, pruned = prune_or_rescue_body_shapes(doc, repl, report["placeholders_missing"])
        report["xml_paras_changed"] = xml_changed
        report["xml_logo_hits"] = xml_logo
        report["xml_paras_pruned"] = pruned
        if xml_changed or xml_logo or pruned:
            report["changed"] = True
    except Exception as xml_error:
        print(f"Error during XML processing: {xml_error}")

    # 6. Shapes in headers/footers
    try:
        hx_changed, hx_logo = rescue_header_footer_shapes(doc, repl, report["placeholders_missing"])
        report["xml_paras_changed"] += hx_changed
        report["xml_logo_hits"] += hx_logo
        if hx_changed or hx_logo:
            report["changed"] = True
    except Exception as hx_error:
        print(f"Error processing header/footer shapes: {hx_error}")

    # 7. Process version control table
    try:
        if not dry:
            print(f"Processing version control table...")
            version_processed = process_version_control_table(doc)
            if version_processed:
                report["version_control_processed"] = 1
                report["changed"] = True
                print(f"Version control table processed successfully")
            else:
                print(f"No version control table found or updated")
        
    except Exception as vc_error:
        print(f"Error processing version control table: {vc_error}")

    # 8. Fallback logo if only in shapes
    if (report["xml_logo_hits"] > 0) and (report["logos_inserted_body"] + report["logos_inserted_headers"] == 0):
        if logo is not None and not dry:
            try:
                first_p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
                # Set right alignment for fallback logo too
                first_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = first_p.add_run()
                run.add_picture(str(logo), width=Mm(width_mm))
                print(f"Fallback logo inserted and right-aligned")
                report["logos_inserted_body"] += 1
                report["changed"] = True
            except Exception:
                pass

    # 9. Save the document
    if not dry:
        try:
            doc.save(str(output_path))
            print(f"Document saved successfully")
        except Exception as save_error:
            print(f"Failed to save document: {save_error}")
            return report

    report["placeholders_found"] = sorted(report["placeholders_found"])
    report["placeholders_missing"] = sorted(report["placeholders_missing"])
    
    print(f"Completed processing: {input_path.name}")
    print(f"Summary: Body logos: {report['logos_inserted_body']}, Headers logos: {report['logos_inserted_headers']}, Changed: {report['changed']}")
    
    return report

# ============================================================================
# FILE MANAGEMENT FUNCTIONS
# ============================================================================

def expand_master(master_src: Path, work: Path) -> Path:
    work.mkdir(parents=True, exist_ok=True)
    dst = work / "Master_Expanded"
    if dst.exists():
        shutil.rmtree(dst)
    if master_src.is_dir():
        shutil.copytree(master_src, dst)
    elif master_src.suffix.lower() == ".zip":
        with zipfile.ZipFile(master_src, "r") as zf:
            zf.extractall(dst)
    else:
        raise ValueError("Master must be a folder or .zip")
    return dst

def copy_selected(master_root: Path, dest_root: Path, services: Optional[List[str]]):
    dest_root.mkdir(parents=True, exist_ok=True)
    wanted = [s.lower() for s in services] if services else None
    for item in master_root.iterdir():
        if item.is_dir() and ((wanted is None) or any(w in item.name.lower() for w in wanted)):
            shutil.copytree(item, dest_root / item.name, dirs_exist_ok=True)
        elif item.is_file() and (wanted is None or any(w in item.parent.name.lower() for w in wanted)):
            shutil.copy2(item, dest_root / item.name)

def walk_docx(root: Path):
    """Walk through directory and find .docx files, excluding temporary files"""
    for p in root.rglob("*.docx"):
        # Skip temporary Word files (start with ~$)
        if p.name.startswith("~$"):
            print(f"Skipping temporary file: {p.name}")
            continue
        # Skip hidden files
        if p.name.startswith("."):
            print(f"Skipping hidden file: {p.name}")
            continue
        yield p

def run_pipeline(master_src: Path, out_dir: Path, data_json: Path, logo: Optional[Path] = None, services_csv: Optional[str] = None, dry_run: bool = False, logo_width_mm: float = 35.0):
    print(f"\nENHANCED PIPELINE STARTED")
    print(f"Logo width received: {logo_width_mm}mm (type: {type(logo_width_mm)})")
    print(f"Master: {master_src}")
    print(f"Output: {out_dir}")
    print(f"Logo: {logo}")
    print(f"Dry run: {dry_run}")
    
    repl = load_replacements(data_json)

    workspace = out_dir.parent / "_work"
    if workspace.exists():
        shutil.rmtree(workspace)
    workspace.mkdir(parents=True, exist_ok=True)

    master_dir = expand_master(master_src, workspace)

    staged = workspace / "staged_client"
    if staged.exists():
        shutil.rmtree(staged)
    staged.mkdir(parents=True, exist_ok=True)
    services = [s.strip() for s in services_csv.split(",")] if services_csv else None
    copy_selected(master_dir, staged, services)

    if out_dir.exists():
        shutil.rmtree(out_dir)
    shutil.copytree(staged, out_dir)

    reports = []
    for docx_path in walk_docx(out_dir):
        print(f"\nProcessing: {docx_path.name}")
        print(f"Using logo width: {logo_width_mm}mm")
        rep = process_docx(docx_path, docx_path, repl, logo=logo if logo and Path(logo).exists() else None, width_mm=logo_width_mm, dry=dry_run)
        reports.append(rep)

    report_path = out_dir.parent / ("dry_run_report.csv" if dry_run else "run_report.csv")
    with open(report_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["file","output","changed","logos_inserted_body","logos_inserted_headers","version_control_processed","xml_logo_hits","xml_paras_changed","xml_paras_pruned","placeholders_found","placeholders_missing"])
        for r in reports:
            w.writerow([
                r["file"], r["output"], r["changed"],
                r["logos_inserted_body"], r["logos_inserted_headers"], r["version_control_processed"],
                r["xml_logo_hits"], r["xml_paras_changed"], r["xml_paras_pruned"],
                "; ".join(r["placeholders_found"]), "; ".join(r["placeholders_missing"]),
            ])
    return report_path

def main():
    ap = argparse.ArgumentParser(description="NDIS templating ENHANCED v2.5 (Smart context-aware logo sizing + Enhanced textbox processing + Fixed version control page breaks)")
    ap.add_argument("--master", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--client", required=True)
    ap.add_argument("--logo")
    ap.add_argument("--services")
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--logo-width-mm", type=float, default=35.0, help="Logo width in millimeters")
    args = ap.parse_args()

    master = Path(args.master)
    out_dir = Path(args.out)
    data = Path(args.client)
    logo = Path(args.logo) if args.logo else None

    report = run_pipeline(master, out_dir, data, logo=logo, services_csv=args.services, dry_run=args.dry_run, logo_width_mm=args.logo_width_mm)
    print(f"Report: {report}")

if __name__ == "__main__":
    main()